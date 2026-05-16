#!/usr/bin/env python3
"""
Convert diary .docx files into JSON import format for this app.

JSON shape produced (TitleCase keys to match exporter/importer):
{
  "entries": [
    { "Date": "YYYY-MM-DD", "Title": "...", "Content": "..." }
  ]
}

Heuristics supported:
- Detect per-entry date from headings or paragraphs that start with a date, e.g.:
  * 2025-11-13
  * 2025/11/13
  * 2025.11.13
  * 2025年11月13日
  * 20251113
- Title is taken from the remainder of the same line after the date (if any).
  Optionally, use the first non-empty paragraph after the date as Title with --title-from-next.
- If a .docx contains no detectable date in content, try to infer from filename (e.g., 2025-11-13.docx, 2025年11月13日.docx, 20251113.docx).
- Multiple entries per .docx are supported (each date heading begins a new entry).

Install dependency:
  pip install python-docx

Usage examples:
  python tools/docx_to_json.py --input path/to/dir --output diary-import.json
  python tools/docx_to_json.py --input path/to/file.docx --title-from-next --updated-now --stdout
"""
from __future__ import annotations
import argparse
import json
import re
import sys
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable, List, Optional, Tuple

try:
    from docx import Document  # type: ignore
except Exception as e:  # pragma: no cover
    print("Missing dependency: python-docx. Install with: pip install python-docx", file=sys.stderr)
    raise

DATE_PATTERNS = [
    # Chinese with spaces: 2025 年 11 月 14 日
    re.compile(r"^\s*(?P<y>\d{4})\s*年\s*(?P<m>\d{1,2})\s*月\s*(?P<d>\d{1,2})\s*日\s*$"),
    # 2025-11-13 / 2025/11/13 / 2025.11.13
    re.compile(r"^\s*(?P<y>\d{4})[-/.](?P<m>\d{1,2})[-/.](?P<d>\d{1,2})\b"),
    # Chinese without spaces: 2025年11月13日
    re.compile(r"^\s*(?P<y>\d{4})年(?P<m>\d{1,2})月(?P<d>\d{1,2})日\s*$"),
    # 20251113
    re.compile(r"^\s*(?P<y>\d{4})(?P<m>\d{2})(?P<d>\d{2})\b"),
]

FILENAME_DATE_PATTERNS = [
    re.compile(r"(?P<y>\d{4})[-_.](?P<m>\d{1,2})[-_.](?P<d>\d{1,2})"),
    re.compile(r"(?P<y>\d{4})年(?P<m>\d{1,2})月(?P<d>\d{1,2})日"),
    re.compile(r"(?P<y>\d{4})(?P<m>\d{2})(?P<d>\d{2})"),
]


def norm_date(y: str, m: str, d: str) -> Optional[str]:
    try:
        dt = datetime(int(y), int(m), int(d))
        return dt.strftime("%Y-%m-%d")
    except Exception:
        return None


def extract_date(s: str) -> Tuple[Optional[str], int]:
    """Return (YYYY-MM-DD, end_index_of_match) if found at start; else (None, -1)."""
    for pat in DATE_PATTERNS:
        m = pat.search(s)
        if m:
            y, mth, d = m.group("y"), m.group("m"), m.group("d")
            nd = norm_date(y, mth, d)
            if nd:
                return nd, m.end()
    return None, -1


def extract_date_from_filename(name: str) -> Optional[str]:
    for pat in FILENAME_DATE_PATTERNS:
        m = pat.search(name)
        if m:
            nd = norm_date(m.group("y"), m.group("m"), m.group("d"))
            if nd:
                return nd
    return None


@dataclass
class Entry:
    date: str
    title: str
    content: str


def parse_docx(doc_path: Path, title_from_next: bool = False) -> List[Entry]:
    doc = Document(str(doc_path))
    entries: List[Entry] = []

    current_date: Optional[str] = None
    current_title: str = ""
    current_content_parts: List[str] = []
    expecting_title_next: bool = False

    def flush():
        nonlocal current_date, current_title, current_content_parts
        if current_date and (current_title or current_content_parts):
            # Preserve blank lines between paragraphs; trim trailing whitespace only
            content = "\n\n".join(current_content_parts).rstrip()
            # Title fallback to empty; app will auto-generate from content when importing
            entries.append(Entry(date=current_date, title=current_title.strip(), content=content))
        current_date = None
        current_title = ""
        current_content_parts = []

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            # preserve blank lines inside content (as separators)
            if current_date is not None:
                current_content_parts.append("")
            continue

        # detect date at the start of line
        d, end_idx = extract_date(text)
        is_heading = False
        try:
            style_name = (para.style.name or "").lower()
            if style_name.startswith("heading") or "标题" in style_name:
                is_heading = True
        except Exception:
            pass

        if d or (is_heading and extract_date(text)[0]):
            if d is None:
                d, end_idx = extract_date(text)
            # start new entry
            flush()
            current_date = d
            remainder = text[end_idx:].strip() if end_idx > -1 else ""
            current_title = remainder
            expecting_title_next = title_from_next and (current_title == "")
            continue

        # capture title from the next paragraph if requested
        if current_date is not None and expecting_title_next:
            current_title = text
            expecting_title_next = False
            continue

        # otherwise accumulate content
        if current_date is None:
            # no date detected yet in document; will try filename later
            current_content_parts.append(text)
        else:
            current_content_parts.append(text)

    # flush last
    flush()

    # If still no entries and we have content, try infer date from filename as a single-entry doc
    if not entries:
        inferred = extract_date_from_filename(doc_path.name)
        if inferred:
            content_all = []
            for p in doc.paragraphs:
                t = (p.text or "").rstrip()
                content_all.append(t)
            content = "\n\n".join(content_all).strip()
            if content:
                entries.append(Entry(date=inferred, title="", content=content))

    return entries


def discover_inputs(path: Path) -> List[Path]:
    if path.is_file() and path.suffix.lower() == ".docx":
        return [path]
    if path.is_dir():
        return sorted([p for p in path.rglob("*.docx") if p.is_file()])
    # glob support
    return sorted([p for p in Path().glob(str(path)) if p.suffix.lower() == ".docx"])  # type: ignore[arg-type]


def main(argv: Optional[Iterable[str]] = None) -> int:
    ap = argparse.ArgumentParser(description="Convert .docx diary files into app-importable JSON")
    ap.add_argument("--input", required=True, help="Path to .docx file, directory, or glob pattern")
    ap.add_argument("--output", help="Output JSON file path (default: diary-import.json)")
    ap.add_argument("--stdout", action="store_true", help="Write JSON to stdout instead of file")
    ap.add_argument("--title-from-next", action="store_true", help="Use first non-empty paragraph after date as Title if same-line title is empty")
    ap.add_argument("--updated-now", action="store_true", help="Include UpdatedAt with current UTC time for each entry")
    args = ap.parse_args(argv)

    input_path = Path(args.input)
    files = discover_inputs(input_path)
    if not files:
        print("No .docx files found for input:", input_path, file=sys.stderr)
        return 2

    all_entries: List[Entry] = []
    for fp in files:
        es = parse_docx(fp, title_from_next=args.title_from_next)
        if not es:
            print(f"Warning: no entries parsed from {fp}", file=sys.stderr)
        all_entries.extend(es)

    # de-dupe by (date, title+content) while preserving order (last wins if duplicates)
    seen = {}
    ordered: List[Entry] = []
    for e in all_entries:
        key = (e.date, e.title.strip(), e.content.strip())
        if key in seen:
            idx = seen[key]
            ordered[idx] = e
        else:
            seen[key] = len(ordered)
            ordered.append(e)

    now_iso = datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ") if args.updated_now else None

    # Build JSON payload with TitleCase keys
    payload_entries = []
    for e in ordered:
        item = {
            "Date": e.date,
            "Title": e.title,
            "Content": e.content,
        }
        if now_iso is not None:
            item["UpdatedAt"] = now_iso
        payload_entries.append(item)

    payload = {"entries": payload_entries}

    out_text = json.dumps(payload, ensure_ascii=False, indent=2)

    if args.stdout or not args.output:
        sys.stdout.write(out_text + "\n")
    else:
        Path(args.output).write_text(out_text, encoding="utf-8")
        print(f"Wrote {args.output} with {len(payload_entries)} entries")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
